#!/usr/bin/env python3
"""Build a concise, searchable DOCX index from OCR caption records."""

from __future__ import annotations

import json
import re
from collections import Counter
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


INPUT_DIR = Path('/private/tmp/screenshots-captions')
OUTPUT = Path('docs/Screenshots视频专家QA汇总_202511-202607.docx')
SOURCE_DIR = '/Users/lukai/Downloads/Screenshots'

TOPICS = {
    'Application & documents': ('apply', 'application', 'document', 'certificate', 'passport', 'cv', 'material', 'submit', 'proposal'),
    'Research fit': ('research', 'researcher', 'field', 'expertise', 'professor', 'industry', 'company'),
    'Program & collaboration': ('project', 'program', 'collaboration', 'cooperation', 'joint', 'partner'),
    'Employment & contract': ('contract', 'position', 'job', 'work', 'salary', 'employ', 'appointment'),
    'China & travel': ('china', 'chinese', 'visit', 'travel', 'visa', 'come to china'),
    'Timeline & results': ('deadline', 'month', 'year', 'result', 'schedule', 'january', 'february', 'march', 'november'),
}
COMMON = {
    'the', 'and', 'you', 'your', 'this', 'that', 'for', 'with', 'from', 'we', 'are', 'is', 'to', 'of', 'in',
    'will', 'can', 'do', 'does', 'did', 'have', 'has', 'please', 'thank', 'yes', 'okay', 'about', 'would',
    'could', 'should', 'our', 'their', 'they', 'there', 'what', 'how', 'when', 'where', 'why', 'which', 'who',
}
QUESTION_RE = re.compile(r'\b(what|why|how|when|where|which|who|can you|could you|would you|do you|does|did|are you|is there|will you)\b', re.I)


def set_font(run, name='Arial Unicode MS', size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    if bold is not None:
        run.bold = bold


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, value in {'top': top, 'start': start, 'bottom': bottom, 'end': end}.items():
        node = tc_mar.find(qn(f'w:{side}'))
        if node is None:
            node = OxmlElement(f'w:{side}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths):
    table.autofit = False
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.first_child_found_in('w:tblW')
    tbl_w.set(qn('w:w'), '9360')
    tbl_w.set(qn('w:type'), 'dxa')
    indent = OxmlElement('w:tblInd')
    indent.set(qn('w:w'), '120')
    indent.set(qn('w:type'), 'dxa')
    table_pr.append(indent)
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    table_pr.append(layout)
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn('w:w'), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn('w:w'), str(width))
            tc_w.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, instr, end])


def apply_styles(doc):
    section = doc.sections[0]
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles['Normal']
    normal.font.name = 'Arial Unicode MS'; normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in [
        ('Heading 1', 16, (46, 116, 181), 18, 10),
        ('Heading 2', 13, (46, 116, 181), 14, 7),
        ('Heading 3', 12, (31, 77, 120), 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = 'Arial Unicode MS'; style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
        style.font.size = Pt(size); style.font.color.rgb = RGBColor(*color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run('Screenshots Expert QA Summary')
    set_font(r, size=9, color=(85, 85, 85), bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run('Searchable review index  |  '); set_font(r, size=9, color=(85, 85, 85))
    page_field(footer)


def text_words(text):
    return re.findall(r"[A-Za-z]{2,}", text.lower())


def meaningful(text):
    words = text_words(text)
    if len(words) < 5 or not (set(words) & COMMON):
        return False
    alpha = sum(char.isalpha() for char in text)
    return alpha / max(1, len(text.replace(' ', ''))) >= 0.45


def topic_names(text):
    lower = text.lower()
    return [name for name, terms in TOPICS.items() if any(term in lower for term in terms)]


def sentence_score(text):
    words = text_words(text)
    return min(len(words) / 10, 4) + 2.5 * len(topic_names(text)) + (3 if '?' in text else 0)


def overlap(a, b):
    aw, bw = set(text_words(a)), set(text_words(b))
    return len(aw & bw) / max(1, len(aw | bw))


def format_time(seconds):
    seconds = int(seconds)
    return f'{seconds // 60:02d}:{seconds % 60:02d}'


def extract_qa(records):
    clean = [item for item in records if meaningful(item['caption'])]
    questions = []
    for index, item in enumerate(clean):
        text = item['caption']
        if '?' not in text and not QUESTION_RE.search(text):
            continue
        if any(overlap(text, prior['question']['caption']) > .60 for prior in questions):
            continue
        answer = None
        for candidate in clean[index + 1:]:
            delta = candidate['time_seconds'] - item['time_seconds']
            if delta > 42:
                break
            if QUESTION_RE.search(candidate['caption']) or '?' in candidate['caption']:
                continue
            if overlap(text, candidate['caption']) < .55:
                answer = candidate
                break
        questions.append({'question': item, 'answer': answer})
    selected = []
    for item in sorted(clean, key=lambda entry: sentence_score(entry['caption']), reverse=True):
        if not topic_names(item['caption']):
            continue
        if any(overlap(item['caption'], chosen['caption']) > .58 for chosen in selected):
            continue
        selected.append(item)
        if len(selected) == 2:
            break
    return questions[:2], selected


def add_label_paragraph(doc, label, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    a = p.add_run(label)
    set_font(a, size=10.5, color=(31, 77, 120), bold=True)
    b = p.add_run(text)
    set_font(b, size=10.5, color=(0, 0, 0))
    b.italic = italic
    return p


def populate_table_cell(cell, text, bold=False, color=(0, 0, 0)):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, size=8.7, color=color, bold=bold)


def main():
    data = []
    for path in sorted(INPUT_DIR.glob('SVID_*.json')):
        entry = json.loads(path.read_text(encoding='utf-8'))
        stem = path.stem
        match = re.match(r'SVID_(\d{8})_(\d{6})', stem)
        if not match:
            continue
        date = datetime.strptime(match.group(1) + match.group(2), '%Y%m%d%H%M%S')
        qas, key_points = extract_qa(entry['caption_records'])
        signal = [item for item in entry['caption_records'] if meaningful(item['caption'])]
        topics = Counter(topic for item in signal for topic in topic_names(item['caption']))
        data.append({
            **entry, 'date': date, 'qas': qas, 'key_points': key_points,
            'meaningful_count': len(signal), 'topics': list(topics),
        })
    data.sort(key=lambda item: item['date'])

    doc = Document(); apply_styles(doc)
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(12); title.paragraph_format.space_after = Pt(4)
    r = title.add_run('Screenshots Expert QA Summary')
    set_font(r, size=24, color=(11, 37, 69), bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    r = subtitle.add_run('134 meeting recordings | 2025-11-05 to 2026-07-31 | Searchable review index')
    set_font(r, size=12, color=(85, 85, 85))

    doc.add_heading('Reading notes', level=1)
    add_label_paragraph(doc, 'Scope: ', f'All {len(data)} MP4 files are covered, totaling {sum(item["duration_seconds"] for item in data) / 3600:.1f} hours. Source folder: {SOURCE_DIR}.')
    add_label_paragraph(doc, 'Method: ', 'Visible English auto-captions were sampled every 6 seconds, OCR-deduplicated, then used to identify question cues and subsequent answer cues.')
    add_label_paragraph(doc, 'Important: ', 'This is a finding and navigation aid, not a verbatim transcript. Re-watch the cited timestamp to confirm important dates, numbers, names, commitments, and contract terms.')

    doc.add_heading('Portfolio overview', level=1)
    topic_video_counts = Counter()
    total_questions = 0
    low_signal = []
    for item in data:
        total_questions += len(item['qas'])
        for topic in item['topics']:
            topic_video_counts[topic] += 1
        if item['meaningful_count'] < 3:
            low_signal.append(item)
    add_label_paragraph(doc, 'Results: ', f'{total_questions} candidate Q&A pairs were detected. {len(low_signal)} recordings had insufficient clear caption evidence and are marked low-signal in the index.')
    topic_table = doc.add_table(rows=1, cols=2)
    topic_table.style = 'Table Grid'; set_table_geometry(topic_table, [4200, 5160])
    for cell, text in zip(topic_table.rows[0].cells, ['Auto-classified topic', 'Videos involved']):
        shade(cell, 'E8EEF5'); populate_table_cell(cell, text, bold=True, color=(31, 77, 120))
    for topic, count in topic_video_counts.most_common():
        cells = topic_table.add_row().cells
        populate_table_cell(cells[0], topic); populate_table_cell(cells[1], str(count))

    doc.add_heading('Video directory', level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'; set_table_geometry(table, [1080, 1200, 960, 1920, 4200])
    headings = ['Date', 'Time', 'Length', 'QA / signal', 'Topic cues']
    for cell, text in zip(table.rows[0].cells, headings):
        shade(cell, 'E8EEF5'); populate_table_cell(cell, text, bold=True, color=(31, 77, 120))
    for item in data:
        row = table.add_row().cells
        populate_table_cell(row[0], item['date'].strftime('%Y-%m-%d'))
        populate_table_cell(row[1], item['date'].strftime('%H:%M'))
        populate_table_cell(row[2], f'{item["duration_seconds"] / 60:.1f}m')
        qa_text = f'{len(item["qas"])} candidate Q&A' if item['meaningful_count'] >= 3 else 'Low-signal / no clear captions'
        populate_table_cell(row[3], qa_text)
        if item['key_points']:
            summary = '；'.join('/'.join(topic_names(point['caption'])) for point in item['key_points'])
        else:
            summary = '—'
        populate_table_cell(row[4], summary or 'General meeting')

    doc.add_page_break()
    doc.add_heading('Per-video Q&A index', level=1)
    current_month = None
    for item in data:
        month = item['date'].strftime('%Y-%m')
        if month != current_month:
            current_month = month
            doc.add_heading(month, level=2)
        heading = doc.add_paragraph(style='Heading 3')
        r = heading.add_run(f"{item['date'].strftime('%Y-%m-%d %H:%M')}  |  {item['video']}  |  {item['duration_seconds'] / 60:.1f} min")
        set_font(r, size=12, color=(31, 77, 120), bold=True)
        if item['meaningful_count'] < 3:
            add_label_paragraph(doc, 'Status: ', 'Low-signal: no sufficiently clear Q&A was detected from visible auto-captions. Re-watch manually if needed.')
            continue
        if item['qas']:
            for number, qa in enumerate(item['qas'], 1):
                q = qa['question']
                add_label_paragraph(doc, f'Question {number} [{format_time(q["time_seconds"])}]: ', q['caption'], italic=True)
                if qa['answer']:
                    a = qa['answer']
                    add_label_paragraph(doc, f'Answer [{format_time(a["time_seconds"])}]: ', a['caption'], italic=True)
                else:
                    add_label_paragraph(doc, 'Answer: ', 'No clear paired answer was detected in the subsequent 42 seconds of sampled captions.')
        else:
            add_label_paragraph(doc, 'Candidate Q&A: ', 'No explicit question was detected; the key statements below are provided for review.')
        if item['key_points']:
            for point in item['key_points']:
                label = f"Key evidence [{format_time(point['time_seconds'])}]: "
                add_label_paragraph(doc, label, point['caption'], italic=True)
        else:
            add_label_paragraph(doc, 'Key evidence: ', 'No high-confidence captions related to application, research, collaboration, employment, travel, or timeline were detected.')

    doc.add_heading('Low-signal recordings', level=1)
    if low_signal:
        for item in low_signal:
            add_label_paragraph(doc, f"{item['date'].strftime('%Y-%m-%d %H:%M')}: ", f"{item['video']} ({item['duration_seconds'] / 60:.1f} min)")
    else:
        add_label_paragraph(doc, 'Result: ', 'None.')
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == '__main__':
    main()
