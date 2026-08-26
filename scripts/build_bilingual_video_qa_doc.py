#!/usr/bin/env python3
"""Create a bilingual copy of the video QA DOCX while preserving its structure."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

CJK_FONT = 'Noto Sans CJK SC'


def iter_all_paragraphs(doc):
    seen = set()

    def emit(paragraphs):
        for paragraph in paragraphs:
            key = id(paragraph._p)
            if key not in seen:
                seen.add(key)
                yield paragraph

    yield from emit(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from emit(cell.paragraphs)
    for section in doc.sections:
        yield from emit(section.header.paragraphs)
        yield from emit(section.footer.paragraphs)


def should_translate(text: str) -> bool:
    text = text.strip()
    if not text or re.fullmatch(r'SVID_\d{8}_\d{6}_1\.mp4', text):
        return False
    if re.fullmatch(r'[\d\-:|. /]+', text):
        return False
    return len(re.findall(r'[A-Za-z]{2,}', text)) >= 2


def append_translation(paragraph, translation: str) -> None:
    if not translation or translation in paragraph.text:
        return
    breaker = paragraph.add_run()
    breaker.add_break()
    run = paragraph.add_run(translation)
    run.font.name = CJK_FONT
    run._element.rPr.rFonts.set(qn('w:ascii'), CJK_FONT)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), CJK_FONT)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), CJK_FONT)
    run._element.rPr.rFonts.set(qn('w:cs'), CJK_FONT)
    run._element.rPr.rFonts.set(qn('w:hint'), 'eastAsia')
    lang = OxmlElement('w:lang')
    lang.set(qn('w:eastAsia'), 'zh-CN')
    run._element.rPr.append(lang)
    if paragraph.style and paragraph.style.name == 'Normal':
        run.font.color.rgb = RGBColor(70, 70, 70)
        run.font.size = Pt(10)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument('source_docx', type=Path)
    parser.add_argument('translation_json', type=Path)
    parser.add_argument('output_docx', type=Path)
    args = parser.parse_args()

    translations = json.loads(args.translation_json.read_text(encoding='utf-8'))
    doc = Document(args.source_docx)
    translated = 0
    missing = []
    for paragraph in iter_all_paragraphs(doc):
        source = paragraph.text.strip()
        if not should_translate(source):
            continue
        translation = translations.get(source)
        if translation:
            append_translation(paragraph, translation)
            translated += 1
        else:
            missing.append(source)
    if missing:
        raise RuntimeError(f'{len(missing)} translation entries are missing; first={missing[0]!r}')
    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output_docx)
    print(f'saved={args.output_docx} translated_paragraphs={translated}')


if __name__ == '__main__':
    main()
